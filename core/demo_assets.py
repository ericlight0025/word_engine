from __future__ import annotations

import csv
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from openpyxl import Workbook


@dataclass
class DemoAssets:
    csv_path: Path
    excel_path: Path
    template_path: Path
    template_paths: list[Path]


DEMO_ROWS = [
    {
        "姓名": "測試甲",
        "日期": "2026-04-05",
        "金額": "50,000",
        "合約編號": "C-001",
        "公司名稱": "示例企業 A",
        "專案名稱": "示例專案 A",
        "付款期限": "2026-04-12",
        "聯絡人": "窗口甲",
    },
    {
        "姓名": "測試乙",
        "日期": "2026-04-06",
        "金額": "30,000",
        "合約編號": "C-002",
        "公司名稱": "示例企業 B",
        "專案名稱": "示例專案 B",
        "付款期限": "2026-04-15",
        "聯絡人": "窗口乙",
    },
    {
        "姓名": "測試丙",
        "日期": "2026-04-07",
        "金額": "",
        "合約編號": "C-003",
        "公司名稱": "示例企業 C",
        "專案名稱": "示例專案 C",
        "付款期限": "2026-04-18",
        "聯絡人": "窗口丙",
    },
]


def build_demo_assets(root: Path) -> DemoAssets:
    demo_dir = root / "assets" / "demo"
    templates_dir = demo_dir / "templates"
    demo_dir.mkdir(parents=True, exist_ok=True)
    templates_dir.mkdir(parents=True, exist_ok=True)

    csv_path = demo_dir / "demo_contracts.csv"
    excel_path = demo_dir / "demo_contracts.xlsx"
    template_specs = [
        ("obsidian_contract_template.docx", "專案合作確認書", "合作公司：{{ 公司名稱 }}"),
        ("obsidian_payment_notice.docx", "付款通知單", "應付款對象：{{ 姓名 }}"),
        ("obsidian_receipt_template.docx", "收款確認書", "收款公司：{{ 公司名稱 }}"),
        ("obsidian_engagement_letter.docx", "顧問委任書", "受委任人：{{ 姓名 }}"),
    ]

    if not csv_path.exists():
        headers = list(DEMO_ROWS[0].keys())
        with csv_path.open("w", encoding="utf-8-sig", newline="") as csv_file:
            writer = csv.DictWriter(csv_file, fieldnames=headers)
            writer.writeheader()
            writer.writerows(DEMO_ROWS)

    if not excel_path.exists():
        workbook = Workbook()
        sheet = workbook.active
        headers = list(DEMO_ROWS[0].keys())
        sheet.title = "合約資料"
        sheet.append(headers)
        for row in DEMO_ROWS:
            sheet.append([row.get(header, "") for header in headers])
        workbook.save(excel_path)

    template_paths: list[Path] = []
    for filename, title, first_line in template_specs:
        template_path = templates_dir / filename
        template_paths.append(template_path)
        if template_path.exists():
            continue

        document = Document()
        section = document.sections[0]
        section.header.paragraphs[0].text = "示範案例｜{{ 公司名稱 }}"
        section.footer.paragraphs[0].text = "窗口：{{ 聯絡人 }}｜付款期限：{{ 付款期限 }}"

        document.add_heading(title, level=0)
        document.add_paragraph("本文件用於示範 Word 版型合併工具的真實商務案例。")
        document.add_paragraph(first_line)
        document.add_paragraph("專案名稱：{{ 專案名稱 }}")
        document.add_paragraph("承辦人：{{ 姓名 }}")
        document.add_paragraph("聯絡窗口：{{ 聯絡人 }}")
        document.add_paragraph("簽署日期：{{ 日期 }}")
        document.add_paragraph("合約編號：{{ 合約編號 }}")
        document.add_paragraph("本期金額：新台幣 {{ 金額 }} 元整")
        document.add_paragraph("付款期限：{{ 付款期限 }}")
        document.add_paragraph("備註：若金額留空，系統仍可產出，但會保留空值供人工複核。")
        document.save(template_path)

    return DemoAssets(
        csv_path=csv_path,
        excel_path=excel_path,
        template_path=template_paths[0],
        template_paths=template_paths,
    )
