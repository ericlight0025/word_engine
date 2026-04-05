from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from openpyxl import Workbook

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from core.excel_reader import read_excel
from core.template_engine import build_tag_statuses, extract_tags, merge_documents


class CoreTests(unittest.TestCase):
    def test_read_excel_skips_blank_rows(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            excel_path = Path(temp_dir) / "sample.xlsx"
            workbook = Workbook()
            sheet = workbook.active
            sheet.append(["姓名", "日期", "金額"])
            sheet.append(["測試甲", "2026-04-05", 50000])
            sheet.append([None, None, None])
            sheet.append(["測試乙", "2026-04-06", 30000])
            workbook.save(excel_path)

            dataset = read_excel(excel_path)

            self.assertEqual(dataset.headers, ["姓名", "日期", "金額"])
            self.assertEqual(len(dataset.rows), 2)
            self.assertEqual(dataset.rows[0]["姓名"], "測試甲")

    def test_extract_tags_reads_document_and_header(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            template_path = Path(temp_dir) / "template.docx"
            document = Document()
            document.add_paragraph("合約人：{{ 姓名 }}")
            header = document.sections[0].header
            header.paragraphs[0].text = "日期：{{日期}}"
            document.save(template_path)

            tags = extract_tags(template_path)

            self.assertEqual(tags, ["姓名", "日期"])

    def test_build_tag_statuses_marks_missing_and_extra(self) -> None:
        statuses = build_tag_statuses(
            tags=["姓名", "日期"],
            headers=["姓名", "金額"],
            sample_row={"姓名": "測試甲", "金額": "50000"},
        )

        status_map = {item.tag: item.status for item in statuses}
        self.assertEqual(status_map["姓名"], "matched")
        self.assertEqual(status_map["日期"], "missing")
        self.assertEqual(status_map["金額"], "extra")

    def test_merge_documents_uses_naming_field_and_fallback(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            template_path = Path(temp_dir) / "contract.docx"
            output_dir = Path(temp_dir) / "out"
            document = Document()
            document.add_paragraph("姓名：{{ 姓名 }}")
            document.add_paragraph("編號：{{ 合約編號 }}")
            document.save(template_path)

            summary = merge_documents(
                template_path=template_path,
                rows=[
                    {"姓名": "測試甲", "合約編號": "C-001"},
                    {"姓名": "測試乙", "合約編號": ""},
                ],
                output_dir=output_dir,
                naming_field="合約編號",
            )

            self.assertEqual(summary.success_count, 2)
            self.assertEqual(summary.warning_count, 1)
            self.assertTrue((output_dir / "C-001.docx").exists())
            self.assertTrue((output_dir / "contract_002.docx").exists())


if __name__ == "__main__":
    unittest.main()
